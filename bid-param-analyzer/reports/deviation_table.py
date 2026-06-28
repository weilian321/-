"""
偏离表生成

将偏离判定结果输出为标准格式的参数偏离表，支持 Word 格式导出。
"""
import os
import uuid

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from config.settings import REPORT_OUTPUT_DIR
from database.models import DeviationResult, RiskLevel, DeviationType


def generate_deviation_table(
    deviation_results: list[DeviationResult],
    param_names: dict[str, str],
    product_params: dict[str, tuple[str, str]],
    output_format: str = "docx",
    task_id: str = "",
) -> str:
    """
    生成标准格式参数偏离表。

    Args:
        deviation_results: 偏离判定结果列表
        param_names: parsed_param_id -> 参数名称 映射
        product_params: match_param_id -> (产品名, 产品值) 映射
        output_format: "docx" 或 "docx" (仅支持 Word)
        task_id: 任务 ID，用于文件命名

    Returns:
        生成的文件路径
    """
    os.makedirs(REPORT_OUTPUT_DIR, exist_ok=True)

    filename = f"deviation_table_{task_id or uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(REPORT_OUTPUT_DIR, filename)

    doc = Document()

    title = doc.add_heading("技术参数偏离表", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["序号", "招标要求", "投标响应", "偏离说明", "偏离状态", "备注"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for idx, result in enumerate(deviation_results, 1):
        row = table.add_row()

        bid_name = param_names.get(result.parsed_param_id, "未知参数")
        bid_value = ""
        prod_value = ""
        if result.match_param_id and result.match_param_id in product_params:
            prod_name, prod_val = product_params[result.match_param_id]
            prod_value = prod_val

        risk_note = ""
        if result.risk_level == RiskLevel.DISQUALIFY.value:
            risk_note = "【废标风险】"
        elif result.risk_level == RiskLevel.SCORE_DEDUCTION.value:
            risk_note = "【扣分项】"

        cells = row.cells
        cells[0].text = str(idx)
        cells[1].text = f"{bid_name}\n要求: {bid_value}"
        cells[2].text = f"响应: {prod_value}"
        cells[3].text = result.explanation
        cells[4].text = result.deviation_type
        cells[5].text = f"{risk_note} {result.suggestion}".strip()

        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

            if result.risk_level == RiskLevel.DISQUALIFY.value:
                _set_cell_shading(cell, "FFD7D7")
            elif result.deviation_type == DeviationType.NEGATIVE.value:
                _set_cell_shading(cell, "FFF3CD")
            elif result.deviation_type == DeviationType.POSITIVE.value:
                _set_cell_shading(cell, "D4EDDA")

    _set_column_widths(table, [0.5, 2.0, 2.0, 2.0, 1.0, 1.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "注：\n"
        "1. 偏离状态分为：正偏离（优于招标要求）、无偏离（满足要求）、负偏离（低于要求）、无法确认（需人工确认）\n"
        "2. 红色标记行为废标级风险项，需重点关注\n"
        "3. 黄色标记行为得分扣分项"
    )

    doc.save(output_path)
    return output_path


def _set_cell_shading(cell, color: str):
    """设置单元格背景色。"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def _set_column_widths(table, widths: list[float]):
    """设置表格列宽（英寸）。"""
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
