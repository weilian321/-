"""
投标分析报告生成

综合决策结果、偏离表、得分数据生成 Word/PDF 格式的投标分析报告。
"""
import os
import uuid

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import REPORT_OUTPUT_DIR


def generate_full_report(
    task_id: str,
    decision: dict,
    advantages: list[dict],
    risks: list[dict],
    suggestions: dict,
    competitive: dict,
    score_summary,
    deviation_results: list,
    param_names: dict[str, str],
    output_format: str = "docx",
) -> str:
    os.makedirs(REPORT_OUTPUT_DIR, exist_ok=True)

    filename = f"analysis_report_{task_id or uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(REPORT_OUTPUT_DIR, filename)

    doc = Document()

    title = doc.add_heading("投标技术分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Decision conclusion
    doc.add_heading("一、投标决策结论", level=1)
    conclusion_text = decision.get("conclusion", "未知")
    confidence = decision.get("confidence", "")
    reason = decision.get("reason", "")
    score_rate = decision.get("score_rate", 0)

    p = doc.add_paragraph()
    run = p.add_run(f"决策结论：{conclusion_text}（置信度：{confidence}）")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(f"技术得分率：{score_rate:.1%}")
    doc.add_paragraph(f"决策依据：{reason}")

    # Score summary
    doc.add_heading("二、技术得分概览", level=1)
    doc.add_paragraph(f"预估总分：{score_summary.total_score:.1f} / {score_summary.max_possible_score:.1f}")
    doc.add_paragraph(f"得分率：{score_summary.score_rate:.1%}")

    # Competitive assessment
    doc.add_heading("三、竞争维度评估", level=1)
    doc.add_paragraph(f"综合评定：{competitive.get('level', '无法评估')}")
    doc.add_paragraph(competitive.get("analysis", ""))
    doc.add_paragraph(
        f"正偏离 {competitive.get('positive_count', 0)} 项 / "
        f"无偏离 {competitive.get('neutral_count', 0)} 项 / "
        f"负偏离 {competitive.get('negative_count', 0)} 项 / "
        f"无法确认 {competitive.get('unconfirmed_count', 0)} 项"
    )

    # Advantages
    doc.add_heading("四、竞争优势项", level=1)
    if advantages:
        for adv in advantages:
            param_name = param_names.get(adv.get("param_id", ""), "未知参数")
            doc.add_paragraph(
                f"{param_name}：{adv.get('explanation', '优于招标要求')}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("暂无显著竞争优势项")

    # Risks
    doc.add_heading("五、风险项清单", level=1)
    if risks:
        for risk in risks:
            param_name = param_names.get(risk.get("param_id", ""), "")
            category = risk.get("category", "")
            level = risk.get("risk_level", "")
            explanation = risk.get("explanation", "")
            text = f"[{category}/{level}]"
            if param_name:
                text += f" {param_name}"
            text += f"：{explanation}"
            doc.add_paragraph(text, style="List Bullet")
    else:
        doc.add_paragraph("未发现明显风险项")

    # Suggestions
    doc.add_heading("六、落地建议", level=1)
    for category, items in suggestions.items():
        if items:
            doc.add_heading(category, level=2)
            for item in items:
                doc.add_paragraph(item, style="List Bullet")

    # Appendix: Deviation table summary
    doc.add_heading("附录：偏离结果摘要", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["参数名称", "偏离状态", "风险等级", "说明"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for d in deviation_results:
        row = table.add_row()
        row.cells[0].text = param_names.get(d.parsed_param_id, "未知")
        row.cells[1].text = d.deviation_type
        row.cells[2].text = d.risk_level
        row.cells[3].text = d.explanation

    doc.save(output_path)
    return output_path
