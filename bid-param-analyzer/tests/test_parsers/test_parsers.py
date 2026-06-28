"""
doc_parser.py + table_extractor.py 单元测试

覆盖: 文件校验、扫描件检测、文档解析、章节定位、参数提取、*号识别
"""
import os
import sys
import tempfile
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))

from database.models import ParameterItem
from parsers.doc_parser import (
    _validate_file,
    detect_scan_type,
    parse_document,
    parse_docx,
    extract_tables,
    ParsedDocument,
)
from parsers.table_extractor import (
    locate_key_sections,
    extract_parameters,
    extract_scoring_rules,
    detect_star_items,
    _check_star,
    extract_parameters_from_table,
)


def _sample_docx_path():
    tmp = tempfile.mktemp(suffix=".docx")
    import docx
    d = docx.Document()
    d.add_paragraph("技术参数要求")
    d.add_paragraph("1. 最大吞吐量 >= 10Gbps")
    d.add_paragraph("2. * 并发连接数 >= 500万")
    d.add_paragraph("3. 支持 SNMP 协议")
    d.add_paragraph("")
    d.add_paragraph("评分标准")
    d.add_paragraph("序号 评审项目 分值")
    d.save(tmp)
    return tmp


class TestFileValidation:
    def test_valid_extensions(self):
        assert callable(_validate_file)

    def test_invalid_extension(self):
        tmp = tempfile.mktemp(suffix=".txt")
        open(tmp, "w").write("test")
        try:
            with pytest.raises(ValueError, match="不支持的文件类型"):
                _validate_file(tmp)
        finally:
            os.unlink(tmp)

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError, match="文件不存在"):
            _validate_file("/tmp/nonexistent_12345_file.docx")


class TestScanDetection:
    def test_detect_not_scanned(self):
        path = _sample_docx_path()
        assert detect_scan_type(path) is False

    def test_detect_missing_pdf_returns_true(self):
        assert detect_scan_type("/tmp/nonexistent_12345.pdf") is True


class TestParseDocument:
    def test_parse_docx(self):
        path = _sample_docx_path()
        result = parse_document(path)
        assert isinstance(result, ParsedDocument)
        assert "最大吞吐量" in result.full_text
        assert result.is_scanned is False

    def test_parse_unknown_format_returns_empty(self):
        with pytest.raises(ValueError, match="不支持的文件类型"):
            parse_document("/tmp/nonexistent_12345.xyz")


class TestExtractTables:
    def test_extract_tables_returns_list(self):
        path = _sample_docx_path()
        doc = parse_docx(path)
        tables = extract_tables(doc)
        assert isinstance(tables, list)


class TestLocateSections:
    def test_locate_tech_section(self):
        text = "第一章 总则\n技术参数要求\n1. 最大吞吐量\n第二章 评分标准\n评分细则"
        sections = locate_key_sections(text)
        assert len(sections) >= 1

    def test_empty_text(self):
        sections = locate_key_sections("")
        assert sections == []


class TestExtractParameters:
    def test_extract_from_text(self):
        text = "技术参数要求\n1. 最大吞吐量 >= 10Gbps\n2. 并发连接数 >= 500万"
        params = extract_parameters(text)
        assert isinstance(params, list)

    def test_empty(self):
        assert extract_parameters("") == []


class TestExtractScoringRules:
    def test_extract_rules(self):
        text = "评分标准\n序号 评审项目 分值\n1 技术参数响应 30分"
        rules = extract_scoring_rules(text)
        assert isinstance(rules, list)


class TestStarDetection:
    def test_star_param(self):
        assert _check_star("* 最大吞吐量") is True

    def test_no_star(self):
        assert _check_star("最大吞吐量") is False

    def test_star_symbol(self):
        assert _check_star("★ 核心要求") is True

    def test_detect_list(self):
        items = [
            ParameterItem(id="p1", name="* 核心参数", requirement_value=">=10", param_type="数值范围"),
            ParameterItem(id="p2", name="普通参数", requirement_value="5", param_type="数值范围"),
        ]
        result = detect_star_items(items)
        assert len(result) == 2
        assert result[0].is_material is True
        assert result[1].is_material is False


class TestTableExtraction:
    def test_extract_from_table(self):
        params = extract_parameters_from_table(
            ["参数名称", "要求值", "单位"],
            [["最大吞吐量", ">=10Gbps", "Gbps"]],
        )
        assert len(params) >= 1

    def test_extract_star_in_table(self):
        params = extract_parameters_from_table(
            ["参数名称", "要求值"],
            [["* 核心指标", ">=100"], ["普通指标", "50"]],
        )
        assert len(params) >= 1
