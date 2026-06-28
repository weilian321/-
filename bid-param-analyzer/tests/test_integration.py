"""
端到端集成测试 (任务 21)

覆盖: 解析/比对/得分/报告全链路
"""
import os
import sys
import tempfile
import uuid
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))



def _insert_parameter(version_id, name, nominal_value, unit="", category=""):
    from database.migrations import get_connection
    conn = get_connection()
    pid = f"par_{uuid.uuid4().hex[:12]}"
    conn.execute(
        "INSERT INTO parameter_records (id, version_id, name, nominal_value, unit, category) VALUES (?, ?, ?, ?, ?, ?)",
        (pid, version_id, name, nominal_value, unit, category),
    )
    conn.commit()
    conn.close()
    return pid


def _setup_product_db():
    import database.repository as repo
    pl = repo.create_product_line("集成测试产品线", "端到端测试")
    v = repo.create_product_version(pl.id, "基准版V1")
    repo.switch_version(pl.id, v.id)
    _insert_parameter(v.id, "最大吞吐量", "12Gbps", unit="Gbps", category="性能")
    _insert_parameter(v.id, "并发连接数", "600万", unit="条", category="性能")
    _insert_parameter(v.id, "新建连接速率", "30万/秒", unit="万/秒", category="性能")
    _insert_parameter(v.id, "电源冗余", "支持", unit="", category="硬件")
    return pl, v


def _create_bid_docx():
    tmp = tempfile.mktemp(suffix=".docx")
    import docx
    d = docx.Document()
    d.add_paragraph("第三章 技术参数要求")
    d.add_paragraph("1. 最大吞吐量: >=10Gbps")
    d.add_paragraph("2. 并发连接数: >=500万")
    d.add_paragraph("3. * 新建连接速率: >=20万/秒")
    d.add_paragraph("4. 电源冗余: 支持")
    d.add_paragraph("5. IPsec吞吐量: >=5Gbps")
    d.add_paragraph("")
    d.add_paragraph("第四章 评分标准")
    d.add_paragraph("1. 技术参数响应 满分25分")
    d.add_paragraph("2. 产品功能完整性 满分10分")
    d.save(tmp)
    return tmp


class TestParseIntegration:
    def test_parse_chain(self):
        pl, v = _setup_product_db()
        bid_path = _create_bid_docx()

        from parsers.doc_parser import parse_document
        doc = parse_document(bid_path)
        assert len(doc.full_text) > 0
        assert doc.page_count > 0

        from parsers.table_extractor import locate_key_sections, extract_parameters
        sections = locate_key_sections(doc.full_text)
        assert len(sections) >= 1

        params = extract_parameters(doc.full_text)
        assert isinstance(params, list)


class TestCompareIntegration:
    def test_compare_chain(self):
        pl, v = _setup_product_db()
        bid_path = _create_bid_docx()

        from parsers.doc_parser import parse_document
        from parsers.table_extractor import locate_key_sections, extract_parameters
        doc = parse_document(bid_path)
        sections = locate_key_sections(doc.full_text)
        bid_params = extract_parameters(doc.full_text)

        import database.repository as repo
        product_params = repo.get_active_params(pl.id)

        from engine.semantic_matcher import match_parameters
        pairs, unmatched = match_parameters(bid_params, product_params)
        assert len(pairs) + len(unmatched) == len(bid_params)

        from engine.deviation_judge import batch_judge
        deviations = batch_judge(pairs, unmatched)
        assert len(deviations) >= len(bid_params)


class TestScoringIntegration:
    def test_scoring_chain(self):
        pl, v = _setup_product_db()
        bid_path = _create_bid_docx()

        from parsers.doc_parser import parse_document
        from parsers.table_extractor import locate_key_sections, extract_parameters
        doc = parse_document(bid_path)
        sections = locate_key_sections(doc.full_text)
        bid_params = extract_parameters(doc.full_text)
        import database.repository as repo
        product_params = repo.get_active_params(pl.id)

        from engine.semantic_matcher import match_parameters
        pairs, unmatched = match_parameters(bid_params, product_params)

        from engine.deviation_judge import batch_judge
        deviations = batch_judge(pairs, unmatched)

        from database.models import ScoringRule
        rules = [
            ScoringRule(id="r1", name="技术响应", max_score=25.0, rule_type="quantitative"),
            ScoringRule(id="r2", name="功能完整性", max_score=10.0, rule_type="qualitative"),
        ]

        param_map = {}
        for i, d in enumerate(deviations):
            param_map[chr(ord("a") + i % 26)] = d.parsed_param_id

        from engine.score_calculator import calculate_scoring
        summary = calculate_scoring(rules, deviations, param_map)
        assert summary.max_possible_score >= 0
        assert 0.0 <= summary.score_rate <= 1.0


class TestReportIntegration:
    def test_report_chain(self):
        pl, v = _setup_product_db()
        bid_path = _create_bid_docx()

        from parsers.doc_parser import parse_document
        from parsers.table_extractor import locate_key_sections, extract_parameters
        doc = parse_document(bid_path)
        sections = locate_key_sections(doc.full_text)
        bid_params = extract_parameters(doc.full_text)
        import database.repository as repo
        product_params = repo.get_active_params(pl.id)

        from engine.semantic_matcher import match_parameters
        pairs, unmatched = match_parameters(bid_params, product_params)

        from engine.deviation_judge import batch_judge
        deviations = batch_judge(pairs, unmatched)

        from database.models import ScoringRule
        rules = [
            ScoringRule(id="r1", name="技术响应", max_score=25.0, rule_type="quantitative"),
            ScoringRule(id="r2", name="功能完整性", max_score=10.0, rule_type="qualitative"),
        ]

        param_map = {}
        for i, d in enumerate(deviations):
            param_map[chr(ord("a") + i % 26)] = d.parsed_param_id

        from engine.score_calculator import calculate_scoring
        summary = calculate_scoring(rules, deviations, param_map)

        from engine.decision_engine import (
            derive_decision,
            generate_advantage_list,
            generate_risk_list,
            generate_suggestions,
            competitive_assessment,
        )
        risk_list = generate_risk_list(deviations, summary)
        decision = derive_decision(summary, risk_list)
        advantages = generate_advantage_list(deviations)
        suggestions = generate_suggestions(risk_list, improvement_items=[])
        competitive = competitive_assessment(deviations, summary)

        assert decision["conclusion"] in ("建议投标", "谨慎投标", "不建议投标")

        from reports.analysis_report import generate_full_report
        from reports.deviation_table import generate_deviation_table

        param_names = {}
        for p in bid_params:
            param_names[p.id] = p.name
        prod_data = {}
        for pp in product_params:
            prod_data[pp.id] = (pp.nominal_value, pp.unit)

        report_path = generate_full_report(
            task_id="integ_test",
            decision=decision,
            advantages=advantages,
            risks=risk_list,
            suggestions={"params": suggestions.get("参数优化方向", []),
                         "evidence": suggestions.get("证明材料准备清单", []),
                         "strategy": suggestions.get("报价策略参考", [])},
            competitive=competitive,
            score_summary=summary,
            deviation_results=deviations,
            param_names=param_names,
        )
        assert os.path.exists(report_path)

        table_path = generate_deviation_table(
            deviation_results=deviations,
            param_names=param_names,
            product_params=prod_data,
            task_id="integ_test",
        )
        assert os.path.exists(table_path)


class TestOrchestratorIntegration:
    def test_orchestrator_full_flow(self):
        pl, v = _setup_product_db()
        bid_path = _create_bid_docx()

        from orchestrator import Orchestrator
        from database.models import TaskStatus

        orch = Orchestrator()
        task_id = orch.create_task(bid_file_path=bid_path, product_line_id=pl.id)
        assert task_id.startswith("task_")

        orch.execute_step(task_id, "parse")
        orch.execute_step(task_id, "compare")
        orch.execute_step(task_id, "score")
        orch.execute_step(task_id, "analyze")

        state = orch.get_task_state(task_id)
        assert state["status"] in (
            TaskStatus.COMPLETED.value,
            TaskStatus.ANALYZING.value,
        )

        context = orch._tasks.get(task_id, {}).get("result", {})
        if "decision" in context:
            assert context["decision"]["conclusion"] in ("建议投标", "谨慎投标", "不建议投标")
