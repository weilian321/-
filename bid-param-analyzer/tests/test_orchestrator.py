"""
orchestrator.py 单元测试

覆盖: 状态机流转、任务创建、断点续算
"""
import os
import sys
import tempfile
import uuid
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))



def _insert_parameter(version_id, name, nominal_value, unit="", category=""):
    from database.migrations import get_connection
    conn = get_connection()
    pid = f"par_{uuid.uuid4().hex[:12]}"
    conn.execute(
        "INSERT INTO parameter_records (id, version_id, name, nominal_value, unit, category) VALUES (?, ?, ?, ?, ?, ?)",
        (pid, version_id, name, nominal_value, unit, category),
    )
    conn.commit()
    conn.close()
    return pid


def _setup_product_db():
    import database.repository as repo
    pl = repo.create_product_line("orch_test_pl", "test")
    v = repo.create_product_version(pl.id, "V1")
    repo.switch_version(pl.id, v.id)
    _insert_parameter(v.id, "最大吞吐量", "10Gbps", unit="Gbps")
    _insert_parameter(v.id, "并发连接数", "500万", unit="条")
    return pl, v


def _create_sample_bid_file():
    tmp = tempfile.mktemp(suffix=".docx")
    import docx
    d = docx.Document()
    d.add_paragraph("技术参数要求")
    d.add_paragraph("1. 最大吞吐量: >=10Gbps")
    d.add_paragraph("2. 并发连接数: >=500万")
    d.add_paragraph("3. * 核心参数: >=100")
    d.add_paragraph("")
    d.add_paragraph("评分标准")
    d.add_paragraph("1. 技术响应 30分")
    d.save(tmp)
    return tmp


class TestTaskCreation:
    def test_create_task(self):
        pl, v = _setup_product_db()
        file_path = _create_sample_bid_file()
        from orchestrator import Orchestrator
        from database.models import TaskStatus

        orch = Orchestrator()
        task_id = orch.create_task(bid_file_path=file_path, product_line_id=pl.id)
        assert task_id.startswith("task_")

        state = orch.get_task_state(task_id)
        assert state["status"] == TaskStatus.PENDING.value

    def test_create_with_product_line_name(self):
        pl, v = _setup_product_db()
        file_path = _create_sample_bid_file()
        from orchestrator import Orchestrator

        orch = Orchestrator()
        task_id = orch.create_task(bid_file_path=file_path, product_line_id=pl.name)
        assert task_id.startswith("task_")

    def test_create_with_invalid_product_line(self):
        from orchestrator import Orchestrator
        import tempfile, os
        fd, tmpfile = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        orch = Orchestrator()
        task_id = orch.create_task(
            bid_file_path=tmpfile,
            product_line_id="nonexistent_line",
        )
        os.unlink(tmpfile)
        assert task_id is not None


class TestStateMachine:
    def test_pending_to_parsing(self):
        pl, v = _setup_product_db()
        file_path = _create_sample_bid_file()
        from orchestrator import Orchestrator
        from database.models import TaskStatus

        orch = Orchestrator()
        task_id = orch.create_task(bid_file_path=file_path, product_line_id=pl.id)

        state = orch.execute_step(task_id, "parse")
        assert state is not None
        current = orch.get_task_state(task_id)
        assert current["status"] in (
            TaskStatus.PARSE_DONE.value,
            TaskStatus.PARSING.value,
        )

    def test_full_flow(self):
        pl, v = _setup_product_db()
        file_path = _create_sample_bid_file()
        from orchestrator import Orchestrator
        from database.models import TaskStatus

        orch = Orchestrator()
        task_id = orch.create_task(bid_file_path=file_path, product_line_id=pl.id)

        orch.execute_step(task_id, "parse")
        orch.execute_step(task_id, "compare")
        orch.execute_step(task_id, "score")
        orch.execute_step(task_id, "analyze")

        state = orch.get_task_state(task_id)
        assert state["status"] in (
            TaskStatus.COMPLETED.value,
            TaskStatus.ANALYZING.value,
        )

        context = orch._tasks.get(task_id, {}).get("result", {})
        if "decision" in context:
            assert "conclusion" in context["decision"]


class TestResumeTask:
    def test_resume_task(self):
        pl, v = _setup_product_db()
        file_path = _create_sample_bid_file()
        from orchestrator import Orchestrator

        orch = Orchestrator()
        task_id = orch.create_task(bid_file_path=file_path, product_line_id=pl.id)
        orch.execute_step(task_id, "parse")
        result = orch.resume_task(task_id)
        assert isinstance(result, bool)


class TestGetTaskState:
    def test_nonexistent_task(self):
        from orchestrator import Orchestrator
        orch = Orchestrator()
        state = orch.get_task_state("nonexistent_task")
        assert "status" in state
