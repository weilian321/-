"""
招标文件解析器

提供 PDF/Word/扫描件的文档解析接口。
运行时通过 MonkeyCode 平台 docparse 与 OCR 能力完成实际解析，
本地仅提供 python-docx 的 Word 解析和通用适配逻辑。
"""
import os
import re
from dataclasses import dataclass, field
from typing import Optional

import docx

from config.settings import (
    SUPPORTED_FILE_TYPES,
    MAX_FILE_SIZE_MB,
    MAX_PAGE_COUNT,
)


@dataclass
class TableData:
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    page: int = 0
    caption: str = ""


@dataclass
class ParsedDocument:
    file_path: str
    full_text: str
    pages: list[str] = field(default_factory=list)
    tables: list[TableData] = field(default_factory=list)
    is_scanned: bool = False
    page_count: int = 0
    metadata: dict = field(default_factory=dict)


def _validate_file(file_path: str) -> None:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_FILE_TYPES:
        raise ValueError(f"不支持的文件类型: {ext}，支持的类型: {SUPPORTED_FILE_TYPES}")

    size_mb = os.path.getsize(file_path) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(
            f"文件大小 {size_mb:.1f}MB 超过上限 {MAX_FILE_SIZE_MB}MB"
        )


def detect_scan_type(file_path: str) -> bool:
    """
    判断文档是否为图片扫描件。

    策略：
    - 检查文件扩展名区分 PDF/Word
    - 对于 PDF：尝试读取文本内容，若文本量极低（<50 字符），判定为扫描件
    - 运行时通过平台 docparse 的 OCR 标志位精确判断
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        try:
            with open(file_path, "rb") as f:
                data = f.read(4096)
            text_signals = data.count(b"BT\n") + data.count(b"/Tj") + data.count(b"TJ")
            if text_signals < 2:
                return True
        except Exception:
            return True
    return False


def _pdf_page_count_from_file(file_path: str) -> int:
    """通过文件二进制数据粗略估算 PDF 页数。"""
    try:
        with open(file_path, "rb") as f:
            content = f.read()
        pages_match = re.findall(rb"/Type\s*/Page[^s]", content)
        return len(pages_match)
    except Exception:
        return 0


def parse_pdf(file_path: str) -> ParsedDocument:
    """
    解析 PDF 招标文件。

    运行时实际委托给 MonkeyCode docparse 完成解析。
    docparse 返回结构包含 full_text 和 tables 数组。
    本地实现提供 fallback 逻辑。
    """
    _validate_file(file_path)
    is_scanned = detect_scan_type(file_path)
    page_count = _pdf_page_count_from_file(file_path)

    if page_count > MAX_PAGE_COUNT:
        raise ValueError(f"PDF 页数 {page_count} 超过上限 {MAX_PAGE_COUNT} 页")

    doc = ParsedDocument(
        file_path=file_path,
        full_text="",
        is_scanned=is_scanned,
        page_count=page_count,
    )

    try:
        import pdfminer.high_level
        doc.full_text = pdfminer.high_level.extract_text(file_path)
    except ImportError:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
                    doc.pages.append(text)
            doc.full_text = "\n\n".join(pages_text)
        except ImportError:
            pass

    return doc


def parse_docx(file_path: str) -> ParsedDocument:
    """
    解析 Word 招标文件。

    使用 python-docx 提取文本、表格、段落结构。
    """
    _validate_file(file_path)

    document = docx.Document(file_path)

    paragraphs_text = []
    tables: list[TableData] = []
    current_page = 1

    for element in document.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = docx.text.paragraph.Paragraph(element, document)
            text = para.text.strip()
            if text:
                paragraphs_text.append(text)

            if "lastRenderedPageBreak" in str(element.xml):
                current_page += 1

        elif tag == "tbl":
            table = docx.table.Table(element, document)
            table_data = []
            headers = []

            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = row_data
                else:
                    table_data.append(row_data)

            if headers or table_data:
                tables.append(TableData(
                    headers=headers,
                    rows=table_data,
                    page=current_page,
                ))

    full_text = "\n\n".join(paragraphs_text)
    page_count = max(current_page, len(full_text) // 3000 + 1)

    if page_count > MAX_PAGE_COUNT:
        raise ValueError(f"文档页数 {page_count} 超过上限 {MAX_PAGE_COUNT} 页")

    return ParsedDocument(
        file_path=file_path,
        full_text=full_text,
        pages=paragraphs_text,
        tables=tables,
        is_scanned=False,
        page_count=page_count,
    )


def extract_tables(parsed_doc: ParsedDocument) -> list[TableData]:
    """
    从解析后的文档中提取所有表格，保持行列结构。

    处理合并单元格：拆分 text 中换行符分隔的内容。
    """
    tables: list[TableData] = []

    for table in parsed_doc.tables:
        normalized_rows = []
        for row in table.rows:
            normalized = []
            for cell in row:
                lines = [l.strip() for l in cell.split("\n") if l.strip()]
                normalized.append("；".join(lines) if len(lines) > 1 else (lines[0] if lines else ""))
            normalized_rows.append(normalized)
        tables.append(TableData(
            headers=table.headers,
            rows=normalized_rows,
            page=table.page,
            caption=table.caption,
        ))

    return tables


def parse_document(file_path: str) -> ParsedDocument:
    """
    统一的文档解析入口，根据文件扩展名自动选择解析器。

    运行时通过 MonkeyCode 平台 docparse 增强解析能力。
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext in (".pdf",):
        return parse_pdf(file_path)
    elif ext in (".doc", ".docx"):
        return parse_docx(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")
